"""=============================================================
c:/1work/Python/djcode/pjtk2/pjtk2/utils/to_docx.py
Created: 03 Nov 2017 13:17:50

DESCRIPTION:

This little script connects to project tracker and extracts the
abstracts for the given year into individual word documents to
facilitate editing and review.  Word documents are saved by project
lead and project code.

05 May 2020 - Updated to use python-docx and new project tracker data
model.  (abstract not comment, and common_lake.abbrev rather than
lake.lake_name)

A. Cottrill
=============================================================

"""

import os
import psycopg2

from docx import Document

# PG_HOST = "142.143.160.56"
PG_HOST = "localhost"
PG_USER = "uglmu"
PG_PWD = "uglmu"
PG_DB = "pjtk2"
# PG_DB = "superior"

OUTDIR = "C:/Users/COTTRILLAD/1work/ScrapBook/ProjectAbstracts"
# OUTDIR = "Y:/Information Resources/Dataset_Utilities/Annual_Report/WordDocs"

YEAR = "2019"
LAKE = "HU"

# LAKE = "Lake Superior"

constr = "host={} dbname={} user={} password={}"

pgconn = psycopg2.connect(constr.format(PG_HOST, PG_DB, PG_USER, PG_PWD))
pgcur = pgconn.cursor()

print("Getting Project Tracker Data...")

sql = """SELECT last_name, PRJ_CD, PRJ_NM, abstract
from pjtk2_project as project
join auth_user on auth_user.id=project.prj_ldr_id
join common_lake as lake on lake.id=project.lake_id
where year= %s and lake.abbrev=%s
order by last_name, prj_cd
"""

pgcur.execute(sql, (YEAR, LAKE))

rs = pgcur.fetchall()

colnames = [x[0] for x in pgcur.description]

print("Making word docs...")

for x in rs:
    record = dict(zip(colnames, x))

    fname = "{}_{}.docx".format(record["last_name"], record["prj_cd"])

    document = Document()
    heading = "{} ({})".format(record["prj_nm"], record["prj_cd"])
    document.add_heading(heading, 0)

    document.add_paragraph("Project abstract:").bold = True

    abstract = record["abstract"]
    for paragraph in abstract.split("\r\n"):
        document.add_paragraph(paragraph)

    document.save(os.path.join(OUTDIR, fname))


pgcur.close()
pgconn.close()
print("Done!!")
